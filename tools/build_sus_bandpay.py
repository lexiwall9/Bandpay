from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output"
OUT.mkdir(exist_ok=True)
DOCX = OUT / "Cuestionario_SUS_Bandpay.docx"
PDF = OUT / "Cuestionario_SUS_Bandpay.pdf"

PURPLE = "5E17EB"
PURPLE_DARK = "3B079D"
PURPLE_LIGHT = "EFE9FE"
INK = "111827"
GRAY = "6B7280"
LINE = "D8DCE6"
PALE = "F8F9FD"
GREEN = "10B981"

ITEMS = [
    "Creo que me gustaría utilizar la aplicación Bandpay con frecuencia.",
    "Considero que la aplicación Bandpay es innecesariamente compleja.",
    "Creo que la aplicación Bandpay es fácil de usar.",
    "Creo que necesitaría el apoyo de una persona con conocimientos técnicos para poder usar Bandpay.",
    "Considero que las diversas funciones de Bandpay están bien integradas.",
    "Creo que hay demasiadas inconsistencias en Bandpay.",
    "Imagino que la mayoría de las personas aprendería a usar Bandpay muy rápidamente.",
    "Encuentro que Bandpay es muy engorrosa de usar.",
    "Me siento muy seguro(a) usando Bandpay.",
    "Necesité aprender muchas cosas antes de poder comenzar a usar Bandpay.",
]

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def borders(cell, color=LINE, size="6"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = tcBorders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            tcBorders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)

def margins(cell, top=70, start=90, bottom=70, end=90):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn("w:tcMar"))
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for side, value in (("top",top),("start",start),("bottom",bottom),("end",end)):
        node = tcMar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value)); node.set(qn("w:type"), "dxa")

def set_table_widths(table, widths_in):
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW"); tblPr.append(tblW)
    total = int(sum(widths_in) * 1440)
    tblW.set(qn("w:w"), str(total)); tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd"); tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "0"); tblInd.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for width in widths_in:
        gc = OxmlElement("w:gridCol"); gc.set(qn("w:w"), str(int(width*1440))); grid.append(gc)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_in):
            cell.width = Inches(width)
            tcW = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tcW.set(qn("w:w"), str(int(width*1440))); tcW.set(qn("w:type"), "dxa")

def set_repeat_header(row):
    trPr = row._tr.get_or_add_trPr(); el = OxmlElement("w:tblHeader"); el.set(qn("w:val"), "true"); trPr.append(el)

def page_num(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página "); run.font.size = Pt(8); run.font.color.rgb = RGBColor.from_string(GRAY)
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])

def text(cell, value, size=9, bold=False, color=INK, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]; p.alignment = align; p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.0
    r = p.add_run(value); r.bold = bold; r.font.name = "Aptos"; r.font.size = Pt(size); r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return p

def add_label_line(doc, pieces):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    for label, width in pieces:
        r = p.add_run(label + " " + "_" * width + "   "); r.font.name="Aptos"; r.font.size=Pt(9); r.font.color.rgb=RGBColor.from_string(INK)

doc = Document()
sec = doc.sections[0]
sec.page_width=Inches(8.5); sec.page_height=Inches(11)
sec.top_margin=Inches(.48); sec.bottom_margin=Inches(.48); sec.left_margin=Inches(.55); sec.right_margin=Inches(.55)
sec.header_distance=Inches(.25); sec.footer_distance=Inches(.25)

styles = doc.styles
normal = styles["Normal"]
normal.font.name="Aptos"; normal.font.size=Pt(9); normal.font.color.rgb=RGBColor.from_string(INK)
normal.paragraph_format.space_after=Pt(4); normal.paragraph_format.line_spacing=1.05
for name, size, before, after in (("Title",22,0,4),("Heading 1",14,10,5),("Heading 2",11,8,4)):
    s=styles[name]; s.font.name="Aptos Display" if name=="Title" else "Aptos"; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(PURPLE_DARK)
    s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True

for section in doc.sections:
    fp=section.footer.paragraphs[0]; page_num(fp)

# Participant page
p=doc.add_paragraph(style="Title"); p.add_run("BANDPAY").font.color.rgb=RGBColor.from_string(PURPLE)
sub=doc.add_paragraph(); sub.paragraph_format.space_after=Pt(6)
r=sub.add_run("Escala de Usabilidad del Sistema (SUS)"); r.bold=True; r.font.size=Pt(13); r.font.color.rgb=RGBColor.from_string(INK)
r=sub.add_run("  |  Cuestionario posterior a la prueba"); r.font.size=Pt(10); r.font.color.rgb=RGBColor.from_string(GRAY)

call=doc.add_table(rows=1, cols=1); set_table_widths(call,[7.4]); c=call.cell(0,0); shade(c,PURPLE_LIGHT); borders(c,PURPLE_LIGHT); margins(c,85,120,85,120)
text(c,"Instrucciones: marque una sola respuesta por afirmación según su experiencia global con Bandpay. Responda de forma espontánea después de completar las tareas de prueba. No hay respuestas correctas o incorrectas.",9,False,INK)

add_label_line(doc, [("Código del participante:",13),("Fecha:",9)])
add_label_line(doc, [("Rol:  ☐ Administrador  ☐ Músico  ☐ Otro:",11),("Dispositivo:",9)])
add_label_line(doc, [("Experiencia previa con Bandpay:  ☐ Ninguna  ☐ Ocasional  ☐ Frecuente",1)])

p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(4)
r=p.add_run("Escala:  "); r.bold=True; r.font.size=Pt(9)
p.add_run("1 = Totalmente en desacuerdo   ·   2 = En desacuerdo   ·   3 = Neutral   ·   4 = De acuerdo   ·   5 = Totalmente de acuerdo").font.size=Pt(8.5)

tbl=doc.add_table(rows=1, cols=6); tbl.alignment=WD_TABLE_ALIGNMENT.LEFT
set_table_widths(tbl,[5.45,.39,.39,.39,.39,.39]); set_repeat_header(tbl.rows[0])
headers=["Afirmación","1","2","3","4","5"]
for i,h in enumerate(headers):
    text(tbl.cell(0,i),h,8.5,True,"FFFFFF",WD_ALIGN_PARAGRAPH.CENTER if i else WD_ALIGN_PARAGRAPH.LEFT); shade(tbl.cell(0,i),PURPLE); borders(tbl.cell(0,i),PURPLE); margins(tbl.cell(0,i),65,75,65,75)
for idx,item in enumerate(ITEMS,1):
    cells=tbl.add_row().cells
    text(cells[0],f"{idx}. {item}",8.5,False,INK)
    for j in range(1,6): text(cells[j],"○",13,False,PURPLE_DARK,WD_ALIGN_PARAGRAPH.CENTER)
    fill="FFFFFF" if idx%2 else PALE
    for c in cells: shade(c,fill); borders(c,LINE); margins(c,55,75,55,75)

p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(5); p.paragraph_format.space_after=Pt(2)
r=p.add_run("Comentario opcional: "); r.bold=True; r.font.size=Pt(8.5)
p.add_run("¿Qué fue lo más fácil o difícil al usar Bandpay?").font.size=Pt(8.5)
for _ in range(2):
    p=doc.add_paragraph("_"*116); p.paragraph_format.space_after=Pt(0); p.runs[0].font.color.rgb=RGBColor.from_string(LINE); p.runs[0].font.size=Pt(8)

# Scoring page
doc.add_page_break()
p=doc.add_paragraph(style="Title"); p.add_run("Hoja de puntuación SUS").font.color.rgb=RGBColor.from_string(PURPLE_DARK)
p=doc.add_paragraph("Uso exclusivo del equipo evaluador · No entregar al participante")
p.runs[0].italic=True; p.runs[0].font.color.rgb=RGBColor.from_string(GRAY); p.paragraph_format.space_after=Pt(8)

doc.add_heading("1. Cálculo", level=1)
steps=[
    "Ítems impares (1, 3, 5, 7 y 9): contribución = respuesta - 1.",
    "Ítems pares (2, 4, 6, 8 y 10): contribución = 5 - respuesta.",
    "Sume las 10 contribuciones y multiplique el total por 2,5. Resultado final: 0 a 100.",
]
for i,s in enumerate(steps,1):
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(.12); p.paragraph_format.space_after=Pt(4)
    rr=p.add_run(f"{i}. "); rr.bold=True; rr.font.color.rgb=RGBColor.from_string(PURPLE)
    p.add_run(s)

score=doc.add_table(rows=1, cols=4); set_table_widths(score,[.55,3.7,1.15,1.95]); score.alignment=WD_TABLE_ALIGNMENT.LEFT
for i,h in enumerate(["Ítem","Regla","Respuesta","Contribución (0-4)"]):
    text(score.cell(0,i),h,8.5,True,"FFFFFF",WD_ALIGN_PARAGRAPH.CENTER if i!=1 else WD_ALIGN_PARAGRAPH.LEFT); shade(score.cell(0,i),PURPLE); borders(score.cell(0,i),PURPLE)
for i in range(1,11):
    row=score.add_row().cells; rule="Respuesta - 1" if i%2 else "5 - respuesta"
    vals=[str(i),rule,"______","______"]
    for j,v in enumerate(vals): text(row[j],v,8.5,False,INK,WD_ALIGN_PARAGRAPH.CENTER if j!=1 else WD_ALIGN_PARAGRAPH.LEFT); shade(row[j],"FFFFFF" if i%2 else PALE); borders(row[j],LINE); margins(row[j],55,85,55,85)

res=doc.add_table(rows=2, cols=4); set_table_widths(res,[1.55,1.65,1.55,2.65]); res.alignment=WD_TABLE_ALIGNMENT.LEFT
labels=["Suma (0-40)","× 2,5","Puntaje SUS","Código participante"]
values=["________","2,5","________ / 100","________________"]
for j,v in enumerate(labels): text(res.cell(0,j),v,8,True,"FFFFFF",WD_ALIGN_PARAGRAPH.CENTER); shade(res.cell(0,j),PURPLE_DARK); borders(res.cell(0,j),PURPLE_DARK)
for j,v in enumerate(values): text(res.cell(1,j),v,10,True,INK,WD_ALIGN_PARAGRAPH.CENTER); shade(res.cell(1,j),PURPLE_LIGHT); borders(res.cell(1,j),PURPLE_LIGHT); margins(res.cell(1,j),90,80,90,80)

doc.add_heading("2. Lectura del resultado", level=1)
p=doc.add_paragraph("El puntaje SUS no es un porcentaje de tareas completadas. Se interpreta comparándolo con resultados de referencia o con mediciones sucesivas del mismo producto. Como referencia general, 68 suele emplearse como promedio aproximado; esta guía es orientativa y no reemplaza el análisis de la muestra.")
p.paragraph_format.space_after=Pt(5)
guide=doc.add_table(rows=2, cols=4); set_table_widths(guide,[1.85,1.85,1.85,1.85])
for j,v in enumerate(["0–50","51–67","68–80","81–100"]): text(guide.cell(0,j),v,9,True,"FFFFFF",WD_ALIGN_PARAGRAPH.CENTER); shade(guide.cell(0,j),PURPLE if j>1 else GRAY); borders(guide.cell(0,j),"FFFFFF")
for j,v in enumerate(["Usabilidad baja","Requiere mejoras","Aceptable / buena","Muy buena / excelente"]): text(guide.cell(1,j),v,8.5,False,INK,WD_ALIGN_PARAGRAPH.CENTER); shade(guide.cell(1,j),PALE if j<2 else PURPLE_LIGHT); borders(guide.cell(1,j),LINE)

doc.add_heading("3. Registro de la sesión", level=1)
add_label_line(doc, [("Tareas completadas:",8),("de:",4),("Duración:",7),("min",0)])
add_label_line(doc, [("Incidencias observadas:",58)])
add_label_line(doc, [("Hallazgo prioritario:",62)])
add_label_line(doc, [("Acción recomendada:",62)])

p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(0)
r=p.add_run("Nota metodológica. "); r.bold=True; r.font.color.rgb=RGBColor.from_string(PURPLE_DARK); r.font.size=Pt(8)
r=p.add_run("Aplicar inmediatamente después de que cada participante complete el mismo conjunto de tareas representativas (por ejemplo: iniciar sesión, revisar un compromiso, consultar un pago y gestionar una notificación). No explicar los ítems durante la respuesta."); r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(GRAY)

doc.core_properties.title="Cuestionario SUS de Bandpay"
doc.core_properties.subject="Instrumento de evaluación de usabilidad listo para aplicación"
doc.core_properties.author="Equipo de evaluación de Bandpay"
doc.save(DOCX)

def make_pdf():
    hx=lambda value: colors.HexColor("#" + value.lstrip("#"))
    pdf = SimpleDocTemplate(str(PDF), pagesize=letter, leftMargin=.55*inch, rightMargin=.55*inch,
                            topMargin=.42*inch, bottomMargin=.42*inch,
                            title="Cuestionario SUS de Bandpay", author="Equipo de evaluación de Bandpay")
    ss=getSampleStyleSheet()
    title=ParagraphStyle("SusTitle", parent=ss["Title"], fontName="Helvetica-Bold", fontSize=21, leading=23, textColor=hx(PURPLE), alignment=0, spaceAfter=3)
    sub=ParagraphStyle("SusSub", parent=ss["Normal"], fontName="Helvetica-Bold", fontSize=11.5, leading=14, textColor=hx(INK), spaceAfter=6)
    body=ParagraphStyle("SusBody", parent=ss["Normal"], fontName="Helvetica", fontSize=8.2, leading=10, textColor=hx(INK), spaceAfter=3)
    small=ParagraphStyle("SusSmall", parent=body, fontSize=7.4, leading=8.8)
    h1=ParagraphStyle("SusH1", parent=ss["Heading1"], fontName="Helvetica-Bold", fontSize=13, leading=15, textColor=hx(PURPLE_DARK), spaceBefore=7, spaceAfter=4)
    center=ParagraphStyle("SusCenter", parent=small, alignment=TA_CENTER)
    story=[]
    story += [Paragraph("BANDPAY",title), Paragraph("Escala de Usabilidad del Sistema (SUS) &nbsp; | &nbsp; Cuestionario posterior a la prueba",sub)]
    call=Table([[Paragraph("<b>Instrucciones:</b> marque una sola respuesta por afirmación según su experiencia global con Bandpay. Responda de forma espontánea después de completar las tareas de prueba. No hay respuestas correctas o incorrectas.",body)]], colWidths=[7.4*inch])
    call.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,-1),hx(PURPLE_LIGHT)),("BOX",(0,0),(-1,-1),.5,hx(PURPLE_LIGHT)),("LEFTPADDING",(0,0),(-1,-1),7),("RIGHTPADDING",(0,0),(-1,-1),7),("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5)]))
    story += [call,Spacer(1,5),Paragraph("Código del participante: ____________________ &nbsp;&nbsp;&nbsp; Fecha: ______________",body),Paragraph("Rol: &nbsp; [ ] Administrador &nbsp; [ ] Músico &nbsp; [ ] Otro: ____________ &nbsp;&nbsp;&nbsp; Dispositivo: ____________",body),Paragraph("Experiencia previa con Bandpay: &nbsp; [ ] Ninguna &nbsp; [ ] Ocasional &nbsp; [ ] Frecuente",body),Spacer(1,2),Paragraph("<b>Escala:</b> 1 = Totalmente en desacuerdo &nbsp; · &nbsp; 2 = En desacuerdo &nbsp; · &nbsp; 3 = Neutral &nbsp; · &nbsp; 4 = De acuerdo &nbsp; · &nbsp; 5 = Totalmente de acuerdo",small)]
    data=[[Paragraph("<b>Afirmación</b>",small)]+[Paragraph(f"<b>{n}</b>",center) for n in range(1,6)]]
    for i,item in enumerate(ITEMS,1): data.append([Paragraph(f"<b>{i}.</b> {item}",body)]+[Paragraph("O",ParagraphStyle(f"c{i}{n}",parent=center,fontSize=9,textColor=hx(PURPLE_DARK))) for n in range(1,6)])
    t=Table(data,colWidths=[5.45*inch]+[.39*inch]*5,repeatRows=1)
    commands=[("BACKGROUND",(0,0),(-1,0),hx(PURPLE)),("TEXTCOLOR",(0,0),(-1,0),colors.white),("GRID",(0,0),(-1,-1),.45,hx(LINE)),("VALIGN",(0,0),(-1,-1),"MIDDLE"),("LEFTPADDING",(0,0),(-1,-1),4),("RIGHTPADDING",(0,0),(-1,-1),4),("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5)]
    for i in range(2,11,2): commands.append(("BACKGROUND",(0,i),(-1,i),hx(PALE)))
    t.setStyle(TableStyle(commands)); story += [t,Spacer(1,4),Paragraph("<b>Comentario opcional:</b> ¿Qué fue lo más fácil o difícil al usar Bandpay?",small),Paragraph("__________________________________________________________________________________________________________________",small),Paragraph("__________________________________________________________________________________________________________________",small)]

    story += [PageBreak(),Paragraph("Hoja de puntuación SUS",title),Paragraph("<i>Uso exclusivo del equipo evaluador · No entregar al participante</i>",body),Paragraph("1. Cálculo",h1)]
    for i,s in enumerate(steps,1): story.append(Paragraph(f"<font color='#{PURPLE}'><b>{i}.</b></font> {s}",body))
    score_data=[[Paragraph(f"<b>{x}</b>",center) for x in ["Ítem","Regla","Respuesta","Contribución (0-4)"]]]
    for i in range(1,11): score_data.append([Paragraph(str(i),center),Paragraph("Respuesta - 1" if i%2 else "5 - respuesta",small),Paragraph("______",center),Paragraph("______",center)])
    st=Table(score_data,colWidths=[.55*inch,3.7*inch,1.15*inch,1.95*inch],repeatRows=1)
    cmds=[("BACKGROUND",(0,0),(-1,0),hx(PURPLE)),("TEXTCOLOR",(0,0),(-1,0),colors.white),("GRID",(0,0),(-1,-1),.45,hx(LINE)),("VALIGN",(0,0),(-1,-1),"MIDDLE"),("LEFTPADDING",(0,0),(-1,-1),5),("RIGHTPADDING",(0,0),(-1,-1),5),("TOPPADDING",(0,0),(-1,-1),3),("BOTTOMPADDING",(0,0),(-1,-1),3)]
    for i in range(2,11,2): cmds.append(("BACKGROUND",(0,i),(-1,i),hx(PALE)))
    st.setStyle(TableStyle(cmds)); story.append(st)
    summary=Table([[Paragraph(f"<b>{v}</b>",center) for v in ["Suma (0-40)","× 2,5","Puntaje SUS","Código participante"]],[Paragraph(f"<b>{v}</b>",center) for v in ["________","2,5","________ / 100","________________"]]],colWidths=[1.55*inch,1.65*inch,1.55*inch,2.65*inch])
    summary.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),hx(PURPLE_DARK)),("TEXTCOLOR",(0,0),(-1,0),colors.white),("BACKGROUND",(0,1),(-1,1),hx(PURPLE_LIGHT)),("GRID",(0,0),(-1,-1),.5,colors.white),("VALIGN",(0,0),(-1,-1),"MIDDLE"),("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5)])); story.append(summary)
    story += [Paragraph("2. Lectura del resultado",h1),Paragraph("El puntaje SUS no es un porcentaje de tareas completadas. Se interpreta comparándolo con resultados de referencia o con mediciones sucesivas del mismo producto. Como referencia general, 68 suele emplearse como promedio aproximado; esta guía es orientativa y no reemplaza el análisis de la muestra.",body)]
    gd=Table([[Paragraph(f"<b>{x}</b>",center) for x in ["0–50","51–67","68–80","81–100"]],[Paragraph(x,center) for x in ["Usabilidad baja","Requiere mejoras","Aceptable / buena","Muy buena / excelente"]]],colWidths=[1.85*inch]*4)
    gd.setStyle(TableStyle([("BACKGROUND",(0,0),(1,0),hx(GRAY)),("BACKGROUND",(2,0),(-1,0),hx(PURPLE)),("TEXTCOLOR",(0,0),(-1,0),colors.white),("BACKGROUND",(0,1),(1,1),hx(PALE)),("BACKGROUND",(2,1),(-1,1),hx(PURPLE_LIGHT)),("GRID",(0,0),(-1,-1),.4,hx(LINE)),("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4)])); story.append(gd)
    story += [Paragraph("3. Registro de la sesión",h1),Paragraph("Tareas completadas: ______ de: ____ &nbsp;&nbsp;&nbsp; Duración: ______ min",body),Paragraph("Incidencias observadas: ______________________________________________________________________________",body),Paragraph("Hallazgo prioritario: __________________________________________________________________________________",body),Paragraph("Acción recomendada: __________________________________________________________________________________",body),Spacer(1,4),Paragraph(f"<font color='#{PURPLE_DARK}'><b>Nota metodológica.</b></font> Aplicar inmediatamente después de que cada participante complete el mismo conjunto de tareas representativas (por ejemplo: iniciar sesión, revisar un compromiso, consultar un pago y gestionar una notificación). No explicar los ítems durante la respuesta.",small)]
    def footer(canvas, doc):
        canvas.saveState(); canvas.setFont("Helvetica",7); canvas.setFillColor(hx(GRAY)); canvas.drawRightString(8.0*inch,.22*inch,f"Página {doc.page}"); canvas.restoreState()
    pdf.build(story,onFirstPage=footer,onLaterPages=footer)

make_pdf()
print(DOCX)
print(PDF)
